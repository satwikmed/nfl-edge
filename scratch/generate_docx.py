import sys
import os
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing via pip...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Plus Jakarta Sans'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(16, 185, 129) # Emerald #10B981
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(6, 182, 212) # Cyan #06B6D4
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(245, 158, 11) # Amber #F59E0B
    return p

def main():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style configurations
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Plus Jakarta Sans'
    font.size = Pt(10)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # Title Page / Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(30)
    title_p.paragraph_format.space_after = Pt(8)
    title_run = title_p.add_run("NFL FRANCHISE TEAM INTELLIGENCE TERMINAL")
    title_run.font.name = 'Plus Jakarta Sans'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 185, 129) # Emerald

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(25)
    sub_run = subtitle_p.add_run("System Showcase, Working Engine Operations, and Advanced Mathematical Modeling Guide")
    sub_run.font.name = 'Plus Jakarta Sans'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500

    doc.add_page_break()

    # Section 1
    add_heading_styled(doc, "1. Introduction to the Platform Showcase", level=1)
    
    p = doc.add_paragraph(
        "The NFL Franchise Team Intelligence Terminal is a professional, high-fidelity quantitative analysis system "
        "designed for sports scientists, head coaches, and front-office scouts. Instead of displaying generic sports statistics "
        "like total yardage or box scores, this interactive platform visualizes how efficiently coaches make in-game decisions, "
        "design play-calling structures, and allocate cap space values."
    )
    p.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph(
        "By utilizing a high-performance SQLite play-by-play database of actual games, the platform extracts raw snap histories, "
        "normalizes them against positional baselines, distributes them across relative bell curves, and showcases the results in a "
        "modern Next.js 15 web application. Every team is graded from A+ down to F across five central columns:"
    )
    p2.paragraph_format.space_after = Pt(10)

    # Section 2
    add_heading_styled(doc, "2. Page-by-Page Interactive Interface Showcase", level=1)
    doc.add_paragraph(
        "Here is a detailed guide on what the application showcases across its four primary navigation sectors, detailing exactly "
        "what the user sees, clicks, and triggers."
    )

    add_heading_styled(doc, "2.1 The Strategic Command Center (Home / Dashboard)", level=2)
    doc.add_paragraph(
        "Upon loading the application, the user is presented with the Strategic Command Center, showcasing an interactive "
        "spatial scatter plot designed to map competitive advantages:"
    )

    sc_items = [
        ("Axis Selectors:", " Located at the top of the dashboard. The user can click dropdown menus to map any of the 5 pillars (or the overall composite score) to either the X-axis or Y-axis. The plot dynamically updates team coordinates, automatically drawing average baseline axes that segment franchises into four strategic quadrants."),
        ("Quadrant 1: Elite Contenders (Top-Right):", " Showcases teams performing above average on both metrics (e.g. green grid items like Green Bay and San Francisco)."),
        ("Quadrant 2: Defensive/Tactical Specialists (Top-Left):", " Features teams with high defensive structure or situational execution but conservative offense (e.g. Baltimore)."),
        ("Quadrant 3: Rebuilding Tier (Bottom-Left):", " Shows franchises currently struggling in both dimensions (e.g. New York Giants, Carolina)."),
        ("Quadrant 4: Offensive Powerhouses (Bottom-Right):", " Highlights high-scoring, schemes that generate positive EPA but lack solid defensive execution."),
        ("Interactive Hover Tooltips:", " Moving the mouse cursor over any team dot on the scatter plot triggers a floating, glassmorphic tooltip box. The tooltip details the franchise name, regular-season record, division, overall composite rating, and their 5 sub-grades."),
        ("Sidebar Quadrant lists:", " Located next to the plot, the sidebar lists all teams mapped into their respective quadrants. Clicking any team abbreviation in this sidebar instantly redirects the user to that team's detail profile page.")
    ]
    for title, desc in sc_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_bold = bp.add_run(title)
        run_bold.font.bold = True
        bp.add_run(desc)

    add_heading_styled(doc, "2.2 The Franchise Head-to-Head Comparison Page (/compare)", level=2)
    doc.add_paragraph(
        "This page allows users to select any two franchises and perform a side-by-side comparative stress test:"
    )

    cmp_items = [
        ("Dual Selector Dropdowns:", " Located at the top, styled with transparent glassmorphism. When the user changes Team A (e.g. Rams) or Team B (e.g. Cowboys), the page triggers client-side fetches for the respective team data files, updating the cards instantly."),
        ("Dual Standings Cards:", " Renders the two team profiles side-by-side, showcasing their regular-season wins/losses, overall composite grades, and absolute ratings for the 5 Pillars."),
        ("The Efficiency Profiles Column:", " Compares granular performance specs: Offensive EPA/play, Defensive EPA/play allowed, base pass rates, and roster average age."),
        ("The Comparative Metrics Column (Center):", " A dashed glass container detailing the absolute variance between the two rosters (e.g. Composite Score Diff in favor of a specific team, Offensive EPA split, and absolute salary cap space used in dollars).")
    ]
    for title, desc in cmp_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_bold = bp.add_run(title)
        run_bold.font.bold = True
        bp.add_run(desc)

    add_heading_styled(doc, "2.3 The Global Rankings Center (/rankings)", level=2)
    doc.add_paragraph(
        "Provides a tabular comparative matrix of all 32 NFL franchises. The headers are interactive: clicking any pillar name "
        "(e.g., 'Play-Calling', 'Roster Cap', 'Game Management') instantly sorts the entire 32-team roster by that metric's percentile "
        "score, allowing users to isolate schematic specialists."
    )

    add_heading_styled(doc, "2.4 The Deep-Dive Franchise Analytics Profile (/team/[id])", level=2)
    doc.add_paragraph(
        "Clicking any team logo redirects to their detail page, which acts as a five-tab command panel showcasing the inner workings "
        "of that franchise:"
    )

    tab_items = [
        ("Interactive Header Info:", " Displays the team abbreviation inside a badge styled with a custom 135-degree linear gradient matching the team's primary and secondary franchise colors, along with division standing, stadium metadata, and regular-season win percentage."),
        ("Tab 1: Overview & Profile:", " Visualizes personnel grouping rates (e.g., 11 vs 12 personnel usage) and pass/run EPA split sliders showing efficiency on passing vs running plays."),
        ("Tab 2: Engine A: Play-Calling:", " Displays the team's situational play tendencies across the four down and distance splits. Highlights their average predictability index and overall play-calling efficiency."),
        ("Tab 3: Engine B: Roster Cap Value:", " Renders an interactive salary table showcasing every active player on the roster, mapping their cap hit, remaining contract years, offensive touches/defensive targets, positional VOR, and absolute Cap Hit Efficiency."),
        ("Tab 4: Engine C: In-Game Decisions:", " Details head coaching efficiency on 4th downs, including correct decision percentages, cumulative Expected Points (EP) left on the table, and wated timeout rates. Includes a list of critical clutch close-game snaps with win probability swings."),
        ("Tab 5: Pitch Spatial Analytics:", " Interactive visual field simulator showcasing a heat-map of spatial efficiency hotspots, mapping exactly where on the field the scheme is generating positive Expected Points Added.")
    ]
    for title, desc in tab_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_bold = bp.add_run(title)
        run_bold.font.bold = True
        bp.add_run(desc)

    doc.add_page_break()

    # Section 3
    add_heading_styled(doc, "3. Behind-the-Scenes: Raw Working of the Analytics Engines", level=1)
    doc.add_paragraph(
        "To understand the calculations behind the numbers showcased in the user interface, we look at the operations of the three "
        "backend data-science engines."
    )

    add_heading_styled(doc, "3.1 Engine A: Play-Calling & Tendency Analysis Working", level=2)
    doc.add_paragraph(
        "The engine calculates play tendencies by querying the play_by_play table inside the nfl_teams.db SQLite database. "
        "It groups snaps into four down and distance situations and checks the ratio of passes to runs. If a coach passes "
        "100% of the time on 3rd & long, or runs 100% of the time on 1st & short, the predictability index increases, signaling a "
        "flaw that opposing defensive coordinators can exploit. The engine averages these variances and combines them with the "
        "offense's average EPA per play to yield the final Play-Calling score."
    )

    add_heading_styled(doc, "3.2 Engine B: Roster VOR & Salary Cap Efficiency Working (Upgraded)", level=2)
    doc.add_paragraph(
        "To calculate roster efficiency, our pipeline applies two highly advanced quantitative criteria:"
    )
    
    vor_details = [
        ("55th Percentile Positional replacement level:", " Instead of checking arbitrary low baselines like the 15th percentile (which inflates VOR scores), we establish a mathematically sound replacement baseline at the 55th percentile of snap volume contribution. This represents a standard freely available player (such as a practice squad call-up, cheap veteran, or late-round pick), ensuring VOR ratings are highly meaningful."),
        ("ActivePlayed Regular-Season Roster Alignment:", " The engine strictly filters rosters and contracts to players who registered at least 1 snap during the active regular-season matchups. This syncs contract cap hits with active contributors on the field, preventing distortions from mid-season roster churn or cuts.")
    ]
    for title, desc in vor_details:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_bold = bp.add_run(title)
        run_bold.font.bold = True
        bp.add_run(desc)

    add_heading_styled(doc, "3.3 Engine C: Game Management & Situational Decisions Working (Upgraded)", level=2)
    doc.add_paragraph(
        "This engine measures coaching intelligence under high-leverage constraints using advanced sports science filters:"
    )
    
    gm_steps = [
        ("4th Down Decision Modeling:", " The script scans all 4th down snaps, identifying the mathematically optimal expected points (EP) choice. It records the delta between the optimal and actual choice, storing it as 'Expected Points Left on the Table'."),
        ("Sophisticated Timeout Classification:", " Rather than flagging timeouts based on superficial scoring margins (which covers valid personnel or injury adjustments), we classify a timeout as 'Wasted' if: 1. It is called in the 1st or 3rd Quarter regardless of score (poor early game clock/personnel management), or 2. It is called in the 2nd Quarter when the score difference is greater than 7 points (non-clutch phase). This targets genuine coaching inefficiencies."),
        ("Win Probability (WP) Leverage-Weighted Clutch EPA:", " Measures team execution in the final 5 minutes of close games (4th Qtr/OT, score margin <= 7 points). Crucially, the engine weights each play's EPA by the Win Probability leverage index: Leverage = wp * (1 - wp) * 4. This scales to a maximum of 1 at 50% WP (high stress) and 0 at decided margins, weighting critical plays (e.g. 45% WP) far more heavily than non-consequential ones.")
    ]
    for title, desc in gm_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.bold = True
        p.add_run(desc)

    add_heading_styled(doc, "3.4 Space and Spatial Smoothing on the Field Pitch Heatmap", level=2)
    doc.add_paragraph(
        "To showcase spatial tendencies cleanly without spiky statistical noise (which occurs when using micro-bins like "
        "5-yard horizontal blocks and left/middle/right splits that aggregate very few plays), the platform groups play snaps into "
        "four large, robust 20-yard horizontal zones: the Red Zone (80-100 yards), the Opponent 40-20 Yard Zone, the Midfield Area (40-60 yards), "
        "and the Own 20-40 Yard Zone. This provides statistically significant sample sizes for every zone, resulting in a beautifully "
        "smoothed and highly meaningful visualization of spatial efficiency."
    )

    doc.add_page_break()

    # Section 4
    add_heading_styled(doc, "4. Behind-the-Scenes: The Data Synchronization Pipeline", level=1)
    doc.add_paragraph(
        "The system maintains a Single Source of Truth (SSOT) architecture, utilizing a SQLite database as the foundation "
        "and synchronizing results to the client web application through an automated ETL pipeline:"
    )

    sync_steps = [
        ("1. DB-Level Querying:", " The recalculate_5_metrics_v3.py script connects to nfl_teams.db and queries play logs, games, and player contracts. Crucially, the script filters results to only include regular-season matchups (game_type = 'REG'). This ensures every team is calculated strictly across standard 17-game schedules, eliminating postseason game inflation (which previously led to impossible records like 17-4)."),
        ("2. Relative Normalization:", " Once raw scores are computed for Play-Calling, 4th Downs, VOR Cap, Defense, and Game Management, the script distributes them across relative bell-curve percentiles. It computes percentile standings and maps them to letter grades (A+ through F)."),
        ("3. JSON Export Layer:", " The script writes the normalized results to the dashboard's public/data directory. It updates teams_summary.json (triggering global standings updates on the dashboard) and all 32 individual team_[id].json files (automatically updating the active detailed charts and roster views).")
    ]
    for title, desc in sync_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(6, 182, 212)
        p.add_run(desc)

    # Section 5
    add_heading_styled(doc, "5. Design Aesthetics & Visual Infrastructure", level=1)
    doc.add_paragraph(
        "To ensure a professional-grade visual experience, the frontend incorporates standard premium design tokens:"
    )

    design_items = [
        ("Typography (Plus Jakarta Sans):", " The entire application uses the high-tech, geometric sans-serif typeface Plus Jakarta Sans, imported directly via next/font/google. This gives all headers, tables, and numerical metrics a futuristic, clean athletic tone."),
        ("Glassmorphic Cards:", " Styled with semi-transparent linear gradients (rgba(255,255,255,0.03)) and subtle borders, allowing the animated WebGL background shader to shine through beautifully without reducing text contrast."),
        ("HSL Grade Color Coding:", " Dynamic text classes map grades to harmonious colors: Vibrant Emerald (#34d399) for A-tier ratings, Electric Teal (#2dd4bf) for B-tier, Warm Amber (#fbbf24) for C-tier, and Coral Red (#f87171) for D/F ratings."),
        ("Responsive SVG Coordinates:", " The scatter plot uses mathematical scaling with a 10% bounds padding to project team ratings into SVG dimensions dynamically based on browser window sizes, ensuring a highly fluid layout on both monitors and mobile screens.")
    ]
    for title, desc in design_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_bold = bp.add_run(title)
        run_bold.font.bold = True
        bp.add_run(desc)

    # Save
    out_path = "/Users/satwikmedipalli/Project2/NFL_Franchise_Team_Intelligence_Terminal_Specification.docx"
    doc.save(out_path)
    print(f"Master Specification docx successfully re-generated and written to: {out_path}")

if __name__ == "__main__":
    main()
